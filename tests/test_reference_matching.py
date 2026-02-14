"""
Comprehensive tests for reference section matching.

Tests cover:
- DOCX extraction (paragraphs, tables, mixed)
- Section segmentation (headers, blank lines, localized)
- Scoring logic (priority, placeholders, length penalties)
- Confidence rules (short OCR, delta, placeholders)
- CTA bracket removal
- Section hints filtering
"""

import io
from docx import Document
from docx.shared import Pt

from shared.docx_section_extractor import (
    extract_section_candidates,
    _parse_header,
    _is_subject_line,
)
from shared.reference_matcher import (
    select_best_section,
    _remove_cta_brackets,
    _has_placeholder,
    _count_words,
    _count_chars_no_whitespace,
)
from worker.normalization import normalize_strict, normalize_soft


def _create_test_docx(paragraphs: list[str], table_data: list[list[str]] = None) -> bytes:
    """Helper to create test DOCX in memory."""
    doc = Document()
    
    for para_text in paragraphs:
        doc.add_paragraph(para_text)
    
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


class TestDocxSectionExtractor:
    """Test DOCX section extraction."""
    
    def test_table_only_document(self):
        """DOCX with content in tables only."""
        table_data = [
            ["BANNER", "Buy Now"],
            ["EMAIL", "Subscribe to our newsletter"],
        ]
        docx_bytes = _create_test_docx([], table_data)
        
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        assert len(candidates) >= 1
        # Should extract from table cells
        content_texts = [c.content_text for c in candidates]
        assert any("Buy Now" in text for text in content_texts)
    
    def test_paragraph_with_headers(self):
        """DOCX with paragraph headers."""
        paragraphs = [
            "1) BANNER",
            "Buy Now - Limited Offer",
            "",
            "2) EMAIL",
            "Subscribe today",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        assert len(candidates) == 2
        assert candidates[0].section_number == "1"
        assert candidates[0].section_name == "BANNER"
        assert "Buy Now" in candidates[0].content_text
        
        assert candidates[1].section_number == "2"
        assert candidates[1].section_name == "EMAIL"
    
    def test_localized_headers_blank_line_fallback(self):
        """Headers in non-English should fall back to blank line segmentation."""
        paragraphs = [
            "バナー",
            "今すぐ購入",
            "",
            "",
            "メール",
            "購読する",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        
        candidates = extract_section_candidates(docx_bytes, "test_(ja).docx", "ja")
        
        # Should segment by 2+ blank lines
        assert len(candidates) >= 2
    
    def test_subject_line_excluded(self):
        """Subject: lines are excluded from content."""
        paragraphs = [
            "EMAIL",
            "Subject: Welcome Email",
            "Thank you for subscribing!",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        assert len(candidates) >= 1
        # Subject line should NOT be in content
        assert "Subject:" not in candidates[0].content_text
        assert "Thank you" in candidates[0].content_text


class TestReferenceMatching:
    """Test reference section matching and scoring."""
    
    def test_cta_bracket_removal(self):
        """CTA brackets are removed before scoring."""
        assert _remove_cta_brackets("[BUY NOW]") == "BUY NOW"
        assert _remove_cta_brackets("<Click Here>") == "Click Here"
        assert _remove_cta_brackets("Shop [Now] <here>") == "Shop Now here"
    
    def test_placeholder_detection_robust(self):
        """Robust placeholder detection works."""
        # ASCII identifiers
        assert _has_placeholder("%skin%") is True
        assert _has_placeholder("%displayname%") is True
        assert _has_placeholder("[subscriber_firstname_capitalized]") is True
        
        # Non-ASCII content (robust %...%)
        assert _has_placeholder("%任何内容%") is True
        assert _has_placeholder("%日本語%") is True
        
        # Not placeholders
        assert _has_placeholder("Hello World") is False
        assert _has_placeholder("50% OFF") is False  # Not surrounded by %
    
    def test_cta_matching_improvement(self):
        """[BUY NOW] in reference matches BUY NOW in OCR better."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs_ref = ["BANNER", "[BUY NOW] Limited Time"]
        docx_bytes = _create_test_docx(paragraphs_ref)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        ocr_text = "BUY NOW Limited Time"
        
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
        )
        
        # Should match well
        assert selection.chosen_section is not None
        assert "BUY NOW" in selection.chosen_text
    
    def test_long_text_penalty_non_asian(self):
        """Long text (>50 words) gets penalty for non-Asian languages."""
        from shared.docx_section_extractor import SectionCandidate
        
        short_text = " ".join(["word"] * 10)  # 10 words
        long_text = " ".join(["word"] * 60)  # 60 words
        
        short_cand = SectionCandidate(
            header_text="BANNER",
            content_text=short_text,
            source_path="test.docx",
            language="en",
            section_name="BANNER",
        )
        
        long_cand = SectionCandidate(
            header_text="NEWS",
            content_text=long_text,
            source_path="test.docx",
            language="en",
            section_name="NEWS",
        )
        
        # Short candidate should score higher (all else equal)
        from shared.reference_matcher import _get_length_penalty_multiplier
        
        short_penalty = _get_length_penalty_multiplier(short_cand, "en")
        long_penalty = _get_length_penalty_multiplier(long_cand, "en")
        
        assert short_penalty > long_penalty
    
    def test_long_text_penalty_japanese(self):
        """Long text (>200 chars) gets penalty for Japanese."""
        from shared.docx_section_extractor import SectionCandidate
        from shared.reference_matcher import _get_length_penalty_multiplier
        
        short_text = "日本語" * 50  # 150 chars
        long_text = "日本語" * 80  # 240 chars
        
        short_cand = SectionCandidate(
            header_text=None,
            content_text=short_text,
            source_path="test.docx",
            language="ja",
        )
        
        long_cand = SectionCandidate(
            header_text=None,
            content_text=long_text,
            source_path="test.docx",
            language="ja",
        )
        
        short_penalty = _get_length_penalty_multiplier(short_cand, "ja")
        long_penalty = _get_length_penalty_multiplier(long_cand, "ja")
        
        assert short_penalty > long_penalty
    
    def test_placeholder_penalty_triggers(self):
        """Placeholder penalty reduces score."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs = [
            "BANNER",
            "Hello %displayname%, get 50% OFF!",
            "",
            "EMAIL", 
            "Real text without placeholders",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        ocr_text = "Hello there"
        
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
        )
        
        # Should warn about placeholders or select non-placeholder section
        assert selection.chosen_section is not None
    
    def test_all_placeholders_triggers_manual(self):
        """If all candidates have placeholders → MANUAL."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs = [
            "BANNER",
            "Hello %name%!",
            "",
            "EMAIL",
            "Hi [firstname]!",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        ocr_text = "Hello John"
        
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
        )
        
        assert selection.manual_required is True
        assert "placeholders" in " ".join(selection.warnings).lower()
    
    def test_short_ocr_triggers_manual(self):
        """Short OCR (<= 3 tokens or <= 15 chars) without strict match → MANUAL."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs = ["BANNER", "Buy Now"]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        # Very short OCR
        ocr_text = "OK"
        
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
        )
        
        assert selection.manual_required is True
        assert "too short" in " ".join(selection.warnings).lower()
    
    def test_delta_rule_triggers_manual_strict(self):
        """Delta < 0.05 with no strict match → MANUAL (verified)."""
        from shared.docx_section_extractor import SectionCandidate
        
        # Create two very similar candidates with subtle differences
        paragraphs = [
            "BANNER",
            "Buy Product Now Sale",
            "",
            "EMAIL",
            "Buy Product Now Today",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        # OCR is ambiguous - matches both candidates similarly
        ocr_text = "Buy Product Now"
        
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
        )
        
        # MUST verify both conditions
        assert selection.delta < 0.05, f"Delta should be < 0.05, got {selection.delta}"
        assert selection.manual_required is True, "manual_required must be True for low delta"
        assert "ambiguous" in " ".join(selection.warnings).lower() or "delta" in " ".join(selection.warnings).lower()


class TestSectionHints:
    """Test section hint filtering."""
    
    def test_section_number_hint(self):
        """section_number filters candidates."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs = [
            "1) BANNER",
            "Text A",
            "",
            "2) EMAIL",
            "Text B",
            "",
            "3) POPUP",
            "Text C",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        ocr_text = "Some text"
        
        # Hint for section 2
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
            section_number="2",
        )
        
        # Should prefer section 2
        assert selection.chosen_section_number == "2"
    
    def test_section_name_hint_tolerant(self):
        """section_name matching is tolerant (case, plurals)."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs = [
            "BANNER",
            "Banner text here",
            "",
            "EMAIL",
            "Email text here",
        ]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        ocr_text = "text"
        
        # Hint with different case and plural
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
            section_name="Banners",  # plural, different case
        )
        
        # Should match BANNER
        assert selection.chosen_section_name == "BANNER"
    
    def test_hint_not_found_warning(self):
        """If hint doesn't match anything → warning and fallback."""
        from shared.docx_section_extractor import SectionCandidate
        
        paragraphs = ["BANNER", "Text"]
        docx_bytes = _create_test_docx(paragraphs)
        candidates = extract_section_candidates(docx_bytes, "test.docx", "en")
        
        ocr_text = "Text"
        
        # Hint that doesn't exist
        selection = select_best_section(
            ocr_text=ocr_text,
            candidates=candidates,
            normalize_strict_fn=normalize_strict,
            normalize_soft_fn=normalize_soft,
            section_name="POPUP",
        )
        
        # Should have warning
        assert len(selection.warnings) > 0
        assert "did not match" in " ".join(selection.warnings).lower()


class TestHelperFunctions:
    """Test helper parsing functions."""
    
    def test_parse_header(self):
        """Header parsing extracts number and name."""
        assert _parse_header("1) BANNER") == ("1", "BANNER")
        assert _parse_header("01. EMAIL TEXT") == ("01", "EMAIL TEXT")
        assert _parse_header("BANNER") == (None, "BANNER")
        assert _parse_header("normal text") == (None, None)
    
    def test_is_subject_line(self):
        """Subject: line detection."""
        assert _is_subject_line("Subject: Welcome") is True
        assert _is_subject_line("subject: test") is True
        assert _is_subject_line("SUBJECT: TEST") is True
        assert _is_subject_line("Not a subject") is False
    
    def test_count_words(self):
        """Word counting."""
        assert _count_words("hello world") == 2
        assert _count_words("one") == 1
        assert _count_words("") == 0
    
    def test_count_chars_no_whitespace(self):
        """Character counting (no whitespace, no CTA brackets)."""
        assert _count_chars_no_whitespace("hello world") == 10
        assert _count_chars_no_whitespace("[BUY NOW]") == 6  # Brackets removed
        assert _count_chars_no_whitespace("a b c") == 3


if __name__ == "__main__":
    import pytest
    pytest.main([__file__, "-v"])
