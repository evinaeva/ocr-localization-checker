"""
DOCX section extraction with support for paragraphs and tables.

Extracts candidate sections from DOCX files for reference matching.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import List, Optional

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


@dataclass
class SectionCandidate:
    """Represents a candidate section extracted from DOCX."""
    
    header_text: Optional[str]  # e.g., "BANNER", "01) EMAIL"
    content_text: str  # The actual text to compare
    source_path: str  # DOCX filename
    language: str  # Language code from filename
    section_number: Optional[str] = None  # e.g., "01", "3"
    section_name: Optional[str] = None  # e.g., "BANNER", "EMAIL"
    
    def __repr__(self) -> str:
        return f"SectionCandidate(name={self.section_name}, num={self.section_number}, chars={len(self.content_text)})"


# High priority section keywords (boost in scoring)
HIGH_PRIORITY_KEYWORDS = {"banner", "pic", "im", "popup"}

# Low priority section keywords (penalty in scoring)
LOW_PRIORITY_KEYWORDS = {"news", "email", "letter", "subject"}

# Common section header patterns (English examples, but works with any language)
SECTION_HEADER_PATTERNS = [
    re.compile(r"^(\d+)[.)]\s*(.+)$", re.IGNORECASE),  # "1) BANNER", "01. EMAIL"
    re.compile(r"^([A-Z][A-Z\s]{2,})$"),  # "BANNER", "EMAIL TEXT"
]


def _extract_text_from_docx(docx_bytes: bytes) -> List[str]:
    """
    Extract all text runs from DOCX in document order, preserving empty lines.
    
    Iterates through document body elements (paragraphs and tables) in order.
    Empty paragraphs are preserved as empty strings.
    
    Returns list of text lines in document order.
    """
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    
    # Iterate through body elements in document order
    for element in doc.element.body:
        # Check if element is a paragraph
        if isinstance(element, CT_P):
            para = element
            # Get paragraph object
            para_obj = None
            for p in doc.paragraphs:
                if p._element == para:
                    para_obj = p
                    break
            
            if para_obj is not None:
                text = para_obj.text.strip()
                # Preserve empty lines by appending empty string
                lines.append(text)
        
        # Check if element is a table
        elif isinstance(element, CT_Tbl):
            table_elem = element
            # Find matching table object
            table_obj = None
            for t in doc.tables:
                if t._element == table_elem:
                    table_obj = t
                    break
            
            if table_obj is not None:
                for row in table_obj.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text.strip()
                            # Preserve empty lines in tables too
                            lines.append(text)
    
    return lines


def _parse_header(line: str) -> tuple[Optional[str], Optional[str]]:
    """
    Parse potential section header line.
    
    Returns (section_number, section_name) or (None, None).
    
    Examples:
        "01) BANNER" -> ("01", "BANNER")
        "BANNER" -> (None, "BANNER")
        "3. Email text" -> ("3", "Email text")
    """
    for pattern in SECTION_HEADER_PATTERNS:
        match = pattern.match(line)
        if match:
            groups = match.groups()
            if len(groups) == 2:
                return groups[0], groups[1]
            elif len(groups) == 1:
                return None, groups[0]
    return None, None


def _is_subject_line(line: str) -> bool:
    """Check if line is email Subject: metadata (should be excluded)."""
    return line.lower().startswith("subject:")


def _segment_by_headers(lines: List[str]) -> List[tuple[Optional[str], List[str]]]:
    """
    Segment lines into sections based on detected headers.
    
    Returns list of (header_text, content_lines) tuples.
    """
    sections = []
    current_header = None
    current_content = []
    
    for line in lines:
        # Skip Subject: lines entirely
        if _is_subject_line(line):
            continue
            
        num, name = _parse_header(line)
        
        if name:  # This is a header
            # Save previous section
            if current_content:
                sections.append((current_header, current_content))
            
            # Start new section
            current_header = line
            current_content = []
        else:
            # Add to current section content
            current_content.append(line)
    
    # Don't forget the last section
    if current_content:
        sections.append((current_header, current_content))
    
    return sections


def _segment_by_blank_lines(lines: List[str]) -> List[tuple[Optional[str], List[str]]]:
    """
    Fallback segmentation: split by 2+ consecutive blank lines.
    
    Returns list of (None, content_lines) tuples (no headers).
    """
    sections = []
    current_block = []
    blank_count = 0
    
    for line in lines:
        # Skip Subject: lines
        if _is_subject_line(line):
            continue
            
        if not line.strip():
            blank_count += 1
        else:
            if blank_count >= 2 and current_block:
                # End of block
                sections.append((None, current_block))
                current_block = []
            
            current_block.append(line)
            blank_count = 0
    
    # Last block
    if current_block:
        sections.append((None, current_block))
    
    return sections


def extract_section_candidates(
    docx_bytes: bytes,
    source_path: str,
    language: str,
) -> List[SectionCandidate]:
    """
    Extract section candidates from DOCX file.
    
    Args:
        docx_bytes: DOCX file content
        source_path: Filename for tracing (should be DOCX filename)
        language: Language code (e.g., "en", "ja", "zh-Hans")
    
    Returns:
        List of SectionCandidate objects
    """
    lines = _extract_text_from_docx(docx_bytes)
    
    if not lines:
        return []
    
    # Try header-based segmentation first
    sections = _segment_by_headers(lines)
    
    # If no headers found, fallback to blank line segmentation
    if not sections or all(header is None for header, _ in sections):
        sections = _segment_by_blank_lines(lines)
    
    # Convert to SectionCandidate objects
    candidates = []
    for header_text, content_lines in sections:
        if not content_lines:
            continue
        
        content_text = "\n".join(content_lines)
        
        # Parse header if present
        section_number = None
        section_name = None
        if header_text:
            section_number, section_name = _parse_header(header_text)
        
        candidate = SectionCandidate(
            header_text=header_text,
            content_text=content_text,
            source_path=source_path,
            language=language,
            section_number=section_number,
            section_name=section_name,
        )
        candidates.append(candidate)
    
    return candidates
